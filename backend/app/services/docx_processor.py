"""
docx_processor.py
Processa arquivos .docx no padrão Medcel e monta estrutura para geração de EPUB.

IMPORTANTE — Track Changes:
Arquivos .docx que passaram por revisão (nome contém "_edit_rev", "_rev" etc.)
frequentemente têm texto marcado como inserido (<w:ins>) ou excluído (<w:del>)
via "Controle de Alterações" do Word. A propriedade padrão `paragraph.text` do
python-docx SÓ lê runs que são filhos diretos do parágrafo — texto dentro de
<w:ins> fica aninhado mais fundo e é silenciosamente ignorado, causando perda
de trechos inteiros de frases.

Este módulo faz a extração direta na árvore XML (via lxml/oxml), garantindo:
- Texto inserido (w:ins) É incluído (pois já foi aceito visualmente no Word)
- Texto excluído (w:del) NÃO é incluído (usa tag w:delText, diferente de w:t)
- Imagens em qualquer profundidade da árvore são detectadas
- Formatação (negrito/itálico) e tamanho de fonte são lidos por run real

Detecta por heurística:
- Título do capítulo
- Autores
- Seções H2 (padrão: "1 INTRODUÇÃO", "2 FISIOPATOLOGIA" etc.)
- Seções H3
- Parágrafos normais
- Imagens com legendas (ordem: legenda -> imagem -> fonte)
- Bloco de referências
- Listas (símbolo literal OU numeração nativa do Word)
"""

import re
import base64
import traceback
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
import structlog

logger = structlog.get_logger()


# ─── DEBUG TEMPORÁRIO — remover depois do diagnóstico ────────────────────────
def _ts() -> str:
    return datetime.now().strftime("%H:%M:%S.%f")[:-3]


def _dbg(tag: str, msg: str) -> None:
    print(f"[DEBUG {_ts()}] [docx_processor:{tag}] {msg}", flush=True)
# ──────────────────────────────────────────────────────────────────────────


# ─── Estruturas de dados ──────────────────────────────────────────────────────

@dataclass
class DocxImage:
    filename: str        # ex: img001.jpg
    data_b64: str        # imagem em base64
    media_type: str      # image/jpeg ou image/png
    caption: str = ""    # legenda (Figura X / Quadro X)
    source: str = ""     # fonte (Fonte: ...)


@dataclass
class DocxBlock:
    """Bloco de conteúdo de uma seção."""
    block_type: str      # "paragraph", "h3", "image", "list_item", "alert", "table"
    content: str = ""    # texto HTML-safe
    image: DocxImage = None
    raw_html: str = ""   # HTML já montado (para tabelas)


@dataclass
class DocxSection:
    """Seção H2 do documento."""
    number: int
    title: str           # ex: "1 INTRODUÇÃO"
    blocks: list[DocxBlock] = field(default_factory=list)


@dataclass
class DocxStructure:
    """Estrutura completa do documento Medcel."""
    title: str
    authors: str
    sections: list[DocxSection] = field(default_factory=list)
    references: list[str] = field(default_factory=list)
    images: list[DocxImage] = field(default_factory=list)
    original_filename: str = ""


# ─── Padrões de detecção ─────────────────────────────────────────────────────

# Aceita tanto "1 INTRODUÇÃO" (maiúsculo) quanto "1 Introdução" (title case) —
# o padrão de formatação varia entre documentos, então não travamos em maiúsculo.
H2_PATTERN = re.compile(
    r"^(\d{1,2})\s+([A-ZÁÉÍÓÚÂÊÎÔÛÃÕÀÜÇ][A-Za-zÁÉÍÓÚÂÊÎÔÛÃÕÀÜÇáéíóúâêîôûãõàüç\s\-\/\(\)]{2,60})$"
)
H3_PATTERN = re.compile(r"^(\d+\.\d+)\s+\S+")
AUTHOR_PATTERN = re.compile(r"[A-Z][a-záéíóú]+\s+[A-Z][a-záéíóú]+")
ALERT_PATTERN = re.compile(r"^(ALERTA|PONTO DE PROVA|ATENÇÃO|IMPORTANTE|DICA|CUIDADO)[:\s]", re.IGNORECASE)
CAPTION_PATTERN = re.compile(r"^(Figura|Quadro|Imagem|Tabela)\s+\d+", re.IGNORECASE)
SOURCE_PATTERN = re.compile(r"^(Fonte|Legenda|Source):", re.IGNORECASE)
BULLET_CHARS = ("▶", "▷", "•", "-", "–")


# ─── Extração robusta de XML (aware de track changes) ────────────────────────

W_DEL = qn("w:del")
W_MOVEFROM = qn("w:moveFrom")
W_INS = qn("w:ins")
W_T = qn("w:t")
W_R = qn("w:r")
W_RPR = qn("w:rPr")
W_B = qn("w:b")
W_I = qn("w:i")
W_SZ = qn("w:sz")
W_VAL = qn("w:val")
W_NUMPR = qn("w:numPr")
W_PPR = qn("w:pPr")


def _is_inside_tag(element, tag) -> bool:
    """Verifica se um elemento está aninhado dentro de uma tag específica (ex: w:del)."""
    parent = element.getparent()
    while parent is not None:
        if parent.tag == tag:
            return True
        parent = parent.getparent()
    return False


def _run_props(r_element) -> dict:
    """Lê negrito, itálico e tamanho de fonte direto do XML do run."""
    bold, italic, size = False, False, None
    rpr = r_element.find(W_RPR)
    if rpr is not None:
        b = rpr.find(W_B)
        i = rpr.find(W_I)
        sz = rpr.find(W_SZ)
        if b is not None and b.get(W_VAL) not in ("0", "false"):
            bold = True
        if i is not None and i.get(W_VAL) not in ("0", "false"):
            italic = True
        if sz is not None and sz.get(W_VAL):
            try:
                size = int(sz.get(W_VAL)) / 2  # meio-pontos -> pontos
            except ValueError:
                _dbg("_run_props", f"EXCEPTION (silenciosa antes) ao parsear w:sz val={sz.get(W_VAL)!r}")
                print(traceback.format_exc(), flush=True)
    return {"bold": bold, "italic": italic, "size": size}


def _get_paragraph_runs(paragraph) -> list[dict]:
    """
    Retorna todos os runs de texto do parágrafo, incluindo os que estão dentro
    de <w:ins> (inserções aceitas/pendentes) e excluindo os que estão dentro
    de <w:del> (exclusões). Cada item: {text, bold, italic, size}.
    """
    runs = []
    for r in paragraph._p.findall(".//" + W_R):
        if _is_inside_tag(r, W_DEL) or _is_inside_tag(r, W_MOVEFROM):
            continue
        text = "".join(t.text or "" for t in r.findall(W_T))
        if not text:
            continue
        props = _run_props(r)
        runs.append({"text": text, **props})
    return runs


def _get_paragraph_text(paragraph) -> str:
    return "".join(r["text"] for r in _get_paragraph_runs(paragraph))


def _paragraph_mark_deleted(paragraph) -> bool:
    """
    Detecta se a MARCA DE FIM DE PARÁGRAFO (pilcrow) foi deletada via Track
    Changes — sinalizado por <w:pPr><w:rPr><w:del/></w:rPr></w:pPr>, sem texto.
    Isso é diferente de deletar texto (w:delText): significa que, ao aceitar as
    alterações, ESTE parágrafo se funde com o próximo (dois <w:p> físicos viram
    um parágrafo lógico só). Se não detectarmos isso, frases ficam cortadas no
    meio quando um trecho é deletado/substituído perto do fim do parágrafo.
    """
    ppr = paragraph._p.find(W_PPR)
    if ppr is None:
        return False
    rpr = ppr.find(W_RPR)
    if rpr is None:
        return False
    return rpr.find(W_DEL) is not None


def _merge_tracked_paragraphs(paragraphs: list) -> list[list]:
    """
    Agrupa parágrafos físicos (python-docx) em parágrafos lógicos, fundindo
    cada parágrafo cuja marca de fim foi deletada com o(s) próximo(s) — simula
    o resultado de "aceitar todas as alterações" no Word.
    """
    groups: list[list] = []
    current: list = []
    for para in paragraphs:
        current.append(para)
        if not _paragraph_mark_deleted(para):
            groups.append(current)
            current = []
    if current:
        groups.append(current)
    return groups


def _group_runs(group: list) -> list[dict]:
    runs = []
    for para in group:
        runs.extend(_get_paragraph_runs(para))
    return runs


def _group_text(group: list) -> str:
    return "".join(r["text"] for r in _group_runs(group))


def _paragraph_has_numbering(paragraph) -> bool:
    """Detecta lista nativa do Word (numPr), independente de texto/símbolo."""
    ppr = paragraph._p.find(W_PPR)
    if ppr is None:
        return False
    return ppr.find(W_NUMPR) is not None


# ─── Helpers de HTML ─────────────────────────────────────────────────────────

def _escape_html(text: str) -> str:
    return (text
            .replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
            .replace('"', "&quot;"))


def _runs_to_html(runs: list[dict]) -> str:
    parts = []
    for r in runs:
        text = _escape_html(r["text"])
        if not text.strip():
            parts.append(text)
            continue
        if r["bold"] and r["italic"]:
            parts.append(f"<b><i>{text}</i></b>")
        elif r["bold"]:
            parts.append(f"<b>{text}</b>")
        elif r["italic"]:
            parts.append(f"<i>{text}</i>")
        else:
            parts.append(text)
    return "".join(parts)


def _is_h2_candidate(text: str, runs: list[dict]) -> bool:
    if H2_PATTERN.match(text.strip()):
        return True
    is_bold = any(r["bold"] for r in runs)
    if is_bold and re.match(r"^\d+\s+\w", text.strip()) and len(text.strip()) < 80:
        return True
    return False


def _is_h3_candidate(text: str) -> bool:
    return bool(H3_PATTERN.match(text.strip()))


def _is_list_item(paragraph, text: str) -> bool:
    if text.startswith(BULLET_CHARS):
        return True
    if paragraph.style and "list" in paragraph.style.name.lower():
        return True
    if _paragraph_has_numbering(paragraph):
        return True
    return False


def _is_reference_block(text: str) -> bool:
    clean = text.strip().upper()
    return clean in ("REFERÊNCIAS", "REFERENCIAS", "REFERÊNCIAS BIBLIOGRÁFICAS", "BIBLIOGRAPHY")


def _strip_bullet_prefix(text: str) -> str:
    """Remove símbolo de bullet literal do início, se houver (evita duplicar com <li>)."""
    stripped = text.lstrip()
    for ch in BULLET_CHARS:
        if stripped.startswith(ch):
            return stripped[len(ch):].strip()
    return text


# ─── Extração de imagens (aware de track changes) ────────────────────────────

def _extract_image_from_paragraph(paragraph, img_counter: list) -> DocxImage | None:
    """Procura uma imagem em qualquer profundidade do parágrafo (inclusive dentro de w:ins)."""
    p_element = paragraph._p
    drawings = (
        p_element.findall(".//" + qn("wp:inline")) +
        p_element.findall(".//" + qn("wp:anchor"))
    )
    for drawing in drawings:
        if _is_inside_tag(drawing, W_DEL) or _is_inside_tag(drawing, W_MOVEFROM):
            continue
        blip = drawing.find(".//" + qn("a:blip"))
        if blip is None:
            continue
        rId = blip.get(qn("r:embed"))
        if not rId:
            continue
        try:
            part = paragraph.part.related_parts[rId]
            img_data = part.blob
            ct = part.content_type
            ext = "jpg" if "jpeg" in ct else "png"
            img_counter[0] += 1
            filename = f"img{img_counter[0]:03d}.{ext}"
            return DocxImage(
                filename=filename,
                data_b64=base64.b64encode(img_data).decode(),
                media_type=ct,
            )
        except Exception as e:
            _dbg("_extract_image_from_paragraph", f"EXCEPTION: {e!r}")
            print(traceback.format_exc(), flush=True)
            logger.warning("image_extraction_failed", error=str(e))
    return None


# ─── Processador principal ───────────────────────────────────────────────────

def analyze_docx(docx_path: str, original_filename: str = "") -> DocxStructure:
    """Analisa o .docx e retorna DocxStructure no padrão Medcel."""
    _dbg("analyze_docx", f"INICIO docx_path={docx_path} original_filename={original_filename}")
    doc = Document(docx_path)
    filename = original_filename or Path(docx_path).name
    _dbg("analyze_docx", f"Document() carregado. paragraphs={len(doc.paragraphs)} tables={len(doc.tables)}")

    # Agrupa parágrafos físicos fundindo os que tiveram a marca de fim deletada
    # via Track Changes (ver _merge_tracked_paragraphs) — cada "group" abaixo é
    # um parágrafo lógico (pode conter 1 ou mais <w:p> físicos).
    groups = _merge_tracked_paragraphs(doc.paragraphs)
    _dbg("analyze_docx", f"_merge_tracked_paragraphs OK -> {len(groups)} grupos logicos")
    images_all: list[DocxImage] = []
    img_counter = [0]

    # ── Passo 1: título e autores ──
    # Título = primeira linha não vazia.
    # Autores = todas as linhas curtas seguintes que parecem nome de pessoa
    # (suporta tanto "Nome • Nome • Nome" em uma linha quanto uma linha por autor).
    title = ""
    author_lines: list[str] = []
    idx = 0
    n = len(groups)

    _dbg("analyze_docx", "Passo 1a: procurando titulo (primeira linha nao vazia)")
    while idx < n:
        text = _group_text(groups[idx]).strip()
        idx += 1
        if text:
            title = text
            break
    _dbg("analyze_docx", f"Passo 1a OK: titulo={title!r} (parou em idx={idx}/{n})")

    _dbg("analyze_docx", "Passo 1b: coletando linhas de autor")
    while idx < n:
        group = groups[idx]
        text = _group_text(group).strip()
        if not text:
            idx += 1
            continue

        runs = _group_runs(group)
        if _is_h2_candidate(text, runs):
            break  # já é a primeira seção, não é mais linha de autor

        # Qualquer linha curta nesta posição (entre o título e a primeira seção)
        # é considerada autor. Não validamos "formato de nome" via regex aqui:
        # tentativas anteriores com regex de nome falhavam em nomes com certos
        # acentos (ex: "João" com ã) ou abreviações ("F."), cortando a lista de
        # autores no meio. A própria posição no documento já garante que é autor.
        if len(text) < 100:
            author_lines.append(text)
            idx += 1
            continue

        break  # parágrafo de corpo normal — encerra a coleta de autores

    authors = " • ".join(author_lines)
    start_idx = idx

    logger.info("docx_header", title=title, authors=authors[:80] if authors else "")

    # ── Passo 2: corpo do documento ──
    sections: list[DocxSection] = []
    references: list[str] = []
    current_section: DocxSection | None = None
    in_references = False
    pending_caption = ""

    body_groups = groups[start_idx:]
    _dbg("analyze_docx", f"Passo 2: iniciando loop do corpo, {len(body_groups)} grupos (start_idx={start_idx})")
    for _gi, group in enumerate(body_groups):
        _dbg("analyze_docx:loop_corpo", f"grupo {_gi + 1}/{len(body_groups)} (paras_fisicos={len(group)})")
        # Imagem — verifica por parágrafo físico do grupo ANTES do texto vazio
        # (parágrafo de imagem às vezes não tem nenhum texto). Parágrafos do
        # grupo que renderam imagem não contribuem texto (mesma regra de antes,
        # só que agora aplicada por parágrafo físico dentro do parágrafo lógico).
        text_paras = []
        for p in group:
            img = _extract_image_from_paragraph(p, img_counter)
            if img:
                # Legenda vem ANTES da imagem no documento -> aplica direto na nova imagem
                if pending_caption:
                    img.caption = pending_caption
                    pending_caption = ""
                images_all.append(img)
                if current_section is None:
                    current_section = DocxSection(number=0, title="")
                    sections.append(current_section)
                current_section.blocks.append(DocxBlock(block_type="image", image=img))
            else:
                text_paras.append(p)

        if not text_paras:
            continue  # grupo era só imagem(ns), sem texto a processar

        runs = _group_runs(text_paras)
        text = "".join(r["text"] for r in runs).strip()
        para = text_paras[0]  # representante do grupo para checagens de estilo/numeração

        if not text:
            continue

        if _is_reference_block(text):
            in_references = True
            continue

        if in_references:
            references.append(text)
            continue

        # Legenda de figura/quadro — guarda para aplicar na PRÓXIMA imagem
        if CAPTION_PATTERN.match(text):
            pending_caption = text
            continue

        # Fonte — vem DEPOIS da imagem -> aplica na ÚLTIMA imagem já adicionada
        if SOURCE_PATTERN.match(text):
            if current_section and current_section.blocks:
                for block in reversed(current_section.blocks):
                    if block.block_type == "image" and block.image:
                        block.image.source = text
                        break
            continue

        if _is_h2_candidate(text, runs):
            section_num = len(sections) + 1
            current_section = DocxSection(number=section_num, title=text)
            sections.append(current_section)
            logger.info("section_detected", number=section_num, title=text[:60])
            continue

        if current_section is None:
            current_section = DocxSection(number=0, title="")
            sections.append(current_section)

        if _is_h3_candidate(text):
            current_section.blocks.append(DocxBlock(block_type="h3", content=text))
            continue

        if _is_list_item(para, text):
            html_content = _runs_to_html(runs)
            if html_content.lstrip().startswith(BULLET_CHARS):
                html_content = _strip_bullet_prefix(html_content)
            current_section.blocks.append(
                DocxBlock(block_type="list_item", content=html_content or _escape_html(_strip_bullet_prefix(text)))
            )
            continue

        if ALERT_PATTERN.match(text):
            current_section.blocks.append(DocxBlock(block_type="alert", content=_runs_to_html(runs)))
            continue

        html_content = _runs_to_html(runs)
        if html_content.strip():
            current_section.blocks.append(DocxBlock(block_type="paragraph", content=html_content))

    _dbg("analyze_docx", f"Passo 2 OK: loop do corpo terminou. secoes={len(sections)} imagens={len(images_all)} referencias={len(references)}")

    # Tabelas
    _dbg("analyze_docx", f"Passo 3: iniciando {len(doc.tables)} tabelas")
    for _ti, table in enumerate(doc.tables):
        _dbg("analyze_docx:loop_tabelas", f"tabela {_ti + 1}/{len(doc.tables)} rows={len(table.rows)} cols={len(table.columns)}")
        html = _table_to_html(table)
        if sections:
            sections[-1].blocks.append(DocxBlock(block_type="table", raw_html=html))
    _dbg("analyze_docx", "Passo 3 OK: tabelas processadas")

    logger.info(
        "docx_analysis_done",
        title=title,
        sections=len(sections),
        images=len(images_all),
        references=len(references),
    )

    _dbg("analyze_docx", f"FIM (sucesso) titulo={title!r}")
    return DocxStructure(
        title=title,
        authors=authors,
        sections=sections,
        references=references,
        images=images_all,
        original_filename=filename,
    )


def _table_to_html(table) -> str:
    _dbg("_table_to_html", f"INICIO rows={len(table.rows)}")
    rows_html = []
    for i, row in enumerate(table.rows):
        cells = []
        for cell in row.cells:
            text = _escape_html(cell.text.strip())
            tag = "th" if i == 0 else "td"
            cells.append(f"<{tag}>{text}</{tag}>")
        rows_html.append(f"<tr>{''.join(cells)}</tr>")
    _dbg("_table_to_html", "FIM")
    return f'<table class="medcel-table">{"".join(rows_html)}</table>'
