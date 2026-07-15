"""
docx_processor_generico.py
Processador DOCX genérico — fallback para arquivos que NÃO seguem o padrão
editorial Medcel (docx_processor.py).

Diferença de propósito:
- docx_processor.py (Medcel): heurística específica — seções numeradas
  ("1 INTRODUÇÃO"), legendas "Figura X"/"Quadro X", bloco de autores,
  bloco de referências. Feito sob medida para o template Medcel.
- Este módulo (Genérico): NÃO assume nenhum padrão editorial. Usa os
  estilos reais do Word (Heading 1, Heading 2...) quando existem; se o
  documento não tiver nenhum heading, todo o conteúdo vira um único
  capítulo corrido, sem tentar adivinhar estrutura.

Reaproveita a mesma extração de XML "segura" contra Controle de Alterações
do Word (w:ins / w:del / w:moveFrom / w:moveTo) usada no processador Medcel,
pois esse problema (perda ou duplicação de texto) pode acontecer em qualquer
.docx, não só nos de padrão Medcel.
"""

import re
import base64
import traceback
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
import structlog

logger = structlog.get_logger()


# ─── DEBUG TEMPORÁRIO — remover depois do diagnóstico ────────────────────────
def _ts() -> str:
    return datetime.now().strftime("%H:%M:%S.%f")[:-3]


def _dbg(tag: str, msg: str) -> None:
    print(f"[DEBUG {_ts()}] [docx_processor_generico:{tag}] {msg}", flush=True)
# ──────────────────────────────────────────────────────────────────────────


# ─── Estruturas de dados ──────────────────────────────────────────────────────

@dataclass
class GenericImage:
    filename: str
    data_b64: str
    media_type: str


@dataclass
class GenericBlock:
    block_type: str   # "heading1", "heading2", "heading3", "paragraph", "list_item", "image", "table"
    content: str = ""
    image: GenericImage = None
    raw_html: str = ""


@dataclass
class GenericStructure:
    title: str
    blocks: list[GenericBlock] = field(default_factory=list)
    images: list[GenericImage] = field(default_factory=list)
    original_filename: str = ""


# ─── Extração robusta de XML (mesma lógica anti track-changes do Medcel) ─────

W_DEL = qn("w:del")
W_MOVEFROM = qn("w:moveFrom")
W_T = qn("w:t")
W_R = qn("w:r")
W_RPR = qn("w:rPr")
W_PPR = qn("w:pPr")
W_B = qn("w:b")
W_I = qn("w:i")
W_VAL = qn("w:val")


def _is_inside_tag(element, tag) -> bool:
    parent = element.getparent()
    while parent is not None:
        if parent.tag == tag:
            return True
        parent = parent.getparent()
    return False


def _run_props(r_element) -> dict:
    bold, italic = False, False
    rpr = r_element.find(W_RPR)
    if rpr is not None:
        b = rpr.find(W_B)
        i = rpr.find(W_I)
        if b is not None and b.get(W_VAL) not in ("0", "false"):
            bold = True
        if i is not None and i.get(W_VAL) not in ("0", "false"):
            italic = True
    return {"bold": bold, "italic": italic}


def _get_paragraph_runs(paragraph) -> list[dict]:
    runs = []
    for r in paragraph._p.findall(".//" + W_R):
        if _is_inside_tag(r, W_DEL) or _is_inside_tag(r, W_MOVEFROM):
            continue
        text = "".join(t.text or "" for t in r.findall(W_T))
        if not text:
            continue
        runs.append({"text": text, **_run_props(r)})
    return runs


def _get_paragraph_text(paragraph) -> str:
    return "".join(r["text"] for r in _get_paragraph_runs(paragraph))


def _paragraph_mark_deleted(paragraph) -> bool:
    """Marca de fim de parágrafo deletada via Track Changes (ver docx_processor.py) —
    quando presente, este parágrafo deve se fundir com o próximo."""
    ppr = paragraph._p.find(W_PPR)
    if ppr is None:
        return False
    rpr = ppr.find(W_RPR)
    if rpr is None:
        return False
    return rpr.find(W_DEL) is not None


def _merge_tracked_paragraphs(paragraphs: list) -> list[list]:
    """Agrupa parágrafos físicos fundindo os que tiveram a marca de fim deletada."""
    groups: list[list] = []
    current: list = []
    for para in paragraphs:
        current.append(para)
        if not _paragraph_mark_deleted(para):
            groups.append(current)
            current = []
    if current:
        groups.append(current)
    return groups


def _group_runs(group: list) -> list[dict]:
    runs = []
    for para in group:
        runs.extend(_get_paragraph_runs(para))
    return runs


def _escape_html(text: str) -> str:
    return (text
            .replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
            .replace('"', "&quot;"))


def _runs_to_html(runs: list[dict]) -> str:
    parts = []
    for r in runs:
        text = _escape_html(r["text"])
        if not text.strip():
            parts.append(text)
            continue
        if r["bold"] and r["italic"]:
            parts.append(f"<b><i>{text}</i></b>")
        elif r["bold"]:
            parts.append(f"<b>{text}</b>")
        elif r["italic"]:
            parts.append(f"<i>{text}</i>")
        else:
            parts.append(text)
    return "".join(parts)


def _heading_level(paragraph) -> int | None:
    """
    Detecta nível de heading pelo ESTILO REAL do Word (Heading 1/2/3, Título 1/2/3
    em versões PT-BR). Ao contrário do processador Medcel, não usamos heurística
    de texto/negrito aqui — só o estilo aplicado de fato no documento.
    """
    if not paragraph.style or not paragraph.style.name:
        return None
    name = paragraph.style.name.lower()
    match = re.search(r"(?:heading|título)\s*(\d)", name)
    if match:
        level = int(match.group(1))
        return level if 1 <= level <= 3 else None
    return None


def _extract_image_from_paragraph(paragraph, img_counter: list) -> GenericImage | None:
    p_element = paragraph._p
    drawings = (
        p_element.findall(".//" + qn("wp:inline")) +
        p_element.findall(".//" + qn("wp:anchor"))
    )
    for drawing in drawings:
        if _is_inside_tag(drawing, W_DEL) or _is_inside_tag(drawing, W_MOVEFROM):
            continue
        blip = drawing.find(".//" + qn("a:blip"))
        if blip is None:
            continue
        rId = blip.get(qn("r:embed"))
        if not rId:
            continue
        try:
            part = paragraph.part.related_parts[rId]
            img_data = part.blob
            ct = part.content_type
            ext = "jpg" if "jpeg" in ct else "png"
            img_counter[0] += 1
            filename = f"img{img_counter[0]:03d}.{ext}"
            return GenericImage(
                filename=filename,
                data_b64=base64.b64encode(img_data).decode(),
                media_type=ct,
            )
        except Exception as e:
            _dbg("_extract_image_from_paragraph", f"EXCEPTION: {e!r}")
            print(traceback.format_exc(), flush=True)
            logger.warning("generic_image_extraction_failed", error=str(e))
    return None


# ─── Processador principal ───────────────────────────────────────────────────

def analyze_docx_generico(docx_path: str, original_filename: str = "") -> GenericStructure:
    """
    Analisa o .docx de forma genérica, sem assumir nenhum padrão editorial.
    Título = primeira linha não vazia OU metadado "title" do arquivo, se houver.
    Estrutura = baseada nos estilos de heading reais do Word, se existirem;
    caso contrário, todo o conteúdo vira um capítulo único corrido.
    """
    _dbg("analyze_docx_generico", f"INICIO docx_path={docx_path} original_filename={original_filename}")
    doc = Document(docx_path)
    filename = original_filename or Path(docx_path).name
    _dbg("analyze_docx_generico", f"Document() carregado. paragraphs={len(doc.paragraphs)} tables={len(doc.tables)}")

    core_title = (doc.core_properties.title or "").strip()
    _dbg("analyze_docx_generico", f"core_title={core_title!r}")

    # Agrupa parágrafos físicos fundindo os que tiveram a marca de fim deletada
    # via Track Changes (ver docx_processor.py::_merge_tracked_paragraphs).
    groups = _merge_tracked_paragraphs(doc.paragraphs)
    _dbg("analyze_docx_generico", f"_merge_tracked_paragraphs OK -> {len(groups)} grupos logicos")
    blocks: list[GenericBlock] = []
    images_all: list[GenericImage] = []
    img_counter = [0]

    title = core_title
    title_taken_from_body = False

    _dbg("analyze_docx_generico", f"iniciando loop principal, {len(groups)} grupos")
    for _gi, group in enumerate(groups):
        _dbg("analyze_docx_generico:loop", f"grupo {_gi + 1}/{len(groups)} (paras_fisicos={len(group)})")
        text_paras = []
        for p in group:
            img = _extract_image_from_paragraph(p, img_counter)
            if img:
                images_all.append(img)
                blocks.append(GenericBlock(block_type="image", image=img))
            else:
                text_paras.append(p)

        if not text_paras:
            continue

        runs = _group_runs(text_paras)
        text = "".join(r["text"] for r in runs).strip()
        para = text_paras[0]
        if not text:
            continue

        # Se não havia título nos metadados, usa a primeira linha não vazia do corpo
        if not title and not title_taken_from_body:
            title = text
            title_taken_from_body = True
            continue

        level = _heading_level(para)
        if level == 1:
            blocks.append(GenericBlock(block_type="heading1", content=text))
            continue
        if level == 2:
            blocks.append(GenericBlock(block_type="heading2", content=text))
            continue
        if level == 3:
            blocks.append(GenericBlock(block_type="heading3", content=text))
            continue

        is_list = _paragraph_has_numbering(para) or (paragraph_style_is_list(para))
        if is_list:
            blocks.append(GenericBlock(block_type="list_item", content=_runs_to_html(runs)))
            continue

        blocks.append(GenericBlock(block_type="paragraph", content=_runs_to_html(runs)))

    _dbg("analyze_docx_generico", f"loop principal OK: blocks={len(blocks)} imagens={len(images_all)}")

    _dbg("analyze_docx_generico", f"iniciando {len(doc.tables)} tabelas")
    for _ti, table in enumerate(doc.tables):
        _dbg("analyze_docx_generico:loop_tabelas", f"tabela {_ti + 1}/{len(doc.tables)} rows={len(table.rows)} cols={len(table.columns)}")
        blocks.append(GenericBlock(block_type="table", raw_html=_table_to_html(table)))
    _dbg("analyze_docx_generico", "tabelas OK")

    if not title:
        title = filename.rsplit(".", 1)[0]

    logger.info(
        "generic_docx_analysis_done",
        title=title,
        blocks=len(blocks),
        images=len(images_all),
    )

    _dbg("analyze_docx_generico", f"FIM (sucesso) titulo={title!r}")
    return GenericStructure(
        title=title,
        blocks=blocks,
        images=images_all,
        original_filename=filename,
    )


def _paragraph_has_numbering(paragraph) -> bool:
    ppr = paragraph._p.find(qn("w:pPr"))
    if ppr is None:
        return False
    return ppr.find(qn("w:numPr")) is not None


def paragraph_style_is_list(paragraph) -> bool:
    return bool(paragraph.style and "list" in paragraph.style.name.lower())


def _table_to_html(table) -> str:
    _dbg("_table_to_html", f"INICIO rows={len(table.rows)}")
    rows_html = []
    for i, row in enumerate(table.rows):
        cells = []
        for cell in row.cells:
            text = _escape_html(cell.text.strip())
            tag = "th" if i == 0 else "td"
            cells.append(f"<{tag}>{text}</{tag}>")
        rows_html.append(f"<tr>{''.join(cells)}</tr>")
    _dbg("_table_to_html", "FIM")
    return f'<table class="generic-table">{"".join(rows_html)}</table>'
